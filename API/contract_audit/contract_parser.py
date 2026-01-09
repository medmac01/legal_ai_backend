"""
Description:
    Contract Parser for extracting text from various document formats.
    Supports PDF, DOCX, and TXT files using PyMuPDF (fitz) and python-docx.

Author: Contract Audit Module
Date: 2025/01/09
"""

import os
import logging
from typing import Optional
from io import BytesIO

import fitz  # PyMuPDF
from docx import Document

logger = logging.getLogger(__name__)


class ContractParser:
    """
    Parser for extracting text content from contract documents.
    
    Supports the following file formats:
    - PDF (using PyMuPDF/fitz)
    - DOCX (using python-docx)
    - TXT (plain text)
    """
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt'}
    
    @classmethod
    def extract_text(cls, file_content: bytes, filename: str) -> str:
        """
        Extract text from a document file.
        
        Args:
            file_content: Binary content of the file
            filename: Name of the file (used to determine file type)
            
        Returns:
            Extracted text content as a string
            
        Raises:
            ValueError: If file type is not supported
            Exception: If extraction fails
        """
        file_extension = os.path.splitext(filename)[1].lower()
        
        if file_extension not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {file_extension}. "
                f"Supported types: {', '.join(cls.SUPPORTED_EXTENSIONS)}"
            )
        
        logger.info(f"Extracting text from {filename} ({file_extension})")
        
        try:
            if file_extension == '.pdf':
                return cls._extract_from_pdf(file_content)
            elif file_extension == '.docx':
                return cls._extract_from_docx(file_content)
            elif file_extension == '.txt':
                return cls._extract_from_txt(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            raise
    
    @staticmethod
    def _extract_from_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF using PyMuPDF (fitz).
        
        Args:
            file_content: Binary content of the PDF file
            
        Returns:
            Extracted text content
        """
        text_parts = []
        
        # Open PDF from bytes
        with fitz.open(stream=file_content, filetype="pdf") as doc:
            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(page_text)
                    logger.debug(f"Extracted text from page {page_num + 1}")
        
        full_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from PDF ({len(doc)} pages)")
        
        return full_text
    
    @staticmethod
    def _extract_from_docx(file_content: bytes) -> str:
        """
        Extract text from DOCX using python-docx.
        
        Args:
            file_content: Binary content of the DOCX file
            
        Returns:
            Extracted text content
        """
        doc = Document(BytesIO(file_content))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        full_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from DOCX")
        
        return full_text
    
    @staticmethod
    def _extract_from_txt(file_content: bytes) -> str:
        """
        Extract text from plain text file.
        
        Args:
            file_content: Binary content of the TXT file
            
        Returns:
            Extracted text content
        """
        # Try UTF-8 first, then fall back to latin-1
        try:
            text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            text = file_content.decode('latin-1')
        
        logger.info(f"Extracted {len(text)} characters from TXT")
        return text
    
    @classmethod
    def extract_text_from_path(cls, file_path: str) -> str:
        """
        Extract text from a file path.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file type is not supported
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        with open(file_path, 'rb') as f:
            file_content = f.read()
        
        filename = os.path.basename(file_path)
        return cls.extract_text(file_content, filename)
