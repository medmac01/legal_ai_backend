"""
Description: 
    NDA Generator Service for generating Non-Disclosure Agreements.
    Uses Jinja2 templates for initial NDA generation and Ollama (OpenAI-compatible API)
    for NDA modification/customization.
    
Author: Adapted from Streamlit app for FastAPI
Date  : 2025/01/07
"""

import os
import logging
from datetime import date
from io import BytesIO
from typing import Optional, Tuple, List, Dict, Any

from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
from docx import Document

from .j2_engine import render_template

# Get a logger instance - use a local logger to avoid circular imports
logger = logging.getLogger(__name__)

# System prompt for the chat/modification phase
CHAT_SYSTEM_PROMPT = """
You are a precise legal contract editor. Your task is to modify the provided Non-Disclosure Agreement (NDA) based *only* on the user's explicit request.

**Your Core Directives:**
1.  **Strict Adherence:** Only make the change requested by the user. Do not add, remove, or alter any other part of the contract.
2.  **Maintain Professional Tone:** Ensure the wording of your modifications is formal, professional, and consistent with the existing legal language of the document.
3.  **Output the Full Document:** After making the requested change, you MUST return the entire, complete, and updated text of the NDA. Do not provide summaries, explanations, or confirmation messages. Your output should be only the contract text itself.
4.  **No Commentary:** Do not include phrases like "Here is the updated version:", "I have made the requested change:", or any other conversational text. The output must be ready to be copied and pasted directly into a legal document.
5.  **If Unclear, Ask:** If the user's request is ambiguous, ask for clarification instead of guessing. For example: "Could you please specify which clause you are referring to for the 'indemnity' change?"

Your role is to act as a silent, precise editing tool. The user provides an instruction, you provide the fully updated document.
"""


class NDAService:
    """
    Service class for generating Non-Disclosure Agreements.
    Uses Jinja2 templates for initial generation and Ollama (OpenAI-compatible API) for modifications.
    """
    
    def __init__(self, base_url: Optional[str] = None, model_id: Optional[str] = None):
        """
        Initialize the NDA Service.
        
        Args:
            base_url: Optional Ollama base URL. If not provided, will try to get from LLM_BASE_URL 
                      environment variable, defaults to http://localhost:11434/v1.
            model_id: Optional model ID. If not provided, will try to get from LLM_MODEL_ID 
                      environment variable, defaults to llama3.1:8b.
        """
        self.base_url = base_url or os.getenv("LLM_BASE_URL", "https://4ff035f7dc21.ngrok-free.app/v1")
        self.model_id = model_id or os.getenv("LLM_MODEL_ID", "gpt-oss:20b")
        self.model = None
        
        # Initialize the LLM for modification features
        try:
            # OpenAI API key is required by the client but ignored by Ollama
            api_key = os.getenv("OPENAI_API_KEY", "ollama")
            
            self.model = ChatOpenAI(
                model_name=self.model_id,
                openai_api_key=api_key,
                base_url=self.base_url,
                temperature=0.2,
                max_tokens=8192,
            )
            
            logger.info(f"NDAService initialized with Ollama API (base_url={self.base_url}, model={self.model_id}) for modifications")
        except Exception as e:
            logger.warning(f"Failed to initialize Ollama LLM: {e}. Modification features will be unavailable.")
            self.model = None
    
    def generate_nda_text(self, user_inputs: dict) -> str:
        """
        Generate NDA text based on user inputs using Jinja2 templates.
        
        Args:
            user_inputs: Dictionary containing NDA parameters:
                - first_party: First party company name
                - first_party_address: Address of the first party
                - first_party_incorporation_state: Incorporation state of the first party
                - first_party_representative: Representative of the first party
                - first_party_registration_number: Registration number of the first party
                - first_party_role: Role of first party ("Receiving Party", "Disclosing Party", "Both (Bilateral)")
                - second_party: Second party company name
                - second_party_address: Address of the second party
                - second_party_incorporation_state: Incorporation state of the second party
                - second_party_representative: Representative of the second party
                - second_party_registration_number: Registration number of the second party
                - purpose: Purpose of disclosure
                - applicable_law: Governing law
                - language: Language of the contract
                - duration: Duration of confidentiality in months
                - date: Effective date of the agreement (ISO format)
                - litigation: Dispute resolution mechanism
        
        Returns:
            Generated NDA text as a string.
        
        Raises:
            ValueError: If required fields are missing.
        """
        # Set default date if not provided
        if "date" not in user_inputs or not user_inputs["date"]:
            user_inputs["date"] = date.today().strftime("%Y-%m-%d")
        
        logger.info(f"Generating NDA for first_party: {user_inputs.get('first_party', 'N/A')}, "
                   f"second_party: {user_inputs.get('second_party', 'N/A')}")
        
        try:
            nda_text = render_template(user_inputs)
            logger.info("NDA text generated successfully using Jinja2 template")
            return nda_text
        except ValueError as e:
            logger.error(f"Validation error generating NDA: {e}")
            raise
        except Exception as e:
            logger.error(f"Error generating NDA text: {e}")
            raise Exception(f"Failed to generate NDA: {str(e)}")
    
    def modify_nda(self, nda_text: str, modification_request: str, 
                   conversation_history: Optional[List[Dict[str, Any]]] = None) -> str:
        """
        Modify an existing NDA using Ollama (OpenAI-compatible API).
        
        Args:
            nda_text: The current NDA text to modify.
            modification_request: User's request for how to modify the NDA.
            conversation_history: Optional list of previous conversation messages
                                  for context. Each message is a dict with 'role' and 'content' keys.
        
        Returns:
            Modified NDA text as a string.
        
        Raises:
            ValueError: If Ollama LLM is not available.
            Exception: If LLM API call fails.
        """
        if not self.model:
            raise ValueError("Ollama LLM is not available. Please check LLM_BASE_URL and LLM_MODEL_ID configuration.")
        
        # Build messages list
        messages = [SystemMessage(content=CHAT_SYSTEM_PROMPT)]
        
        if conversation_history:
            # Convert conversation history to LangChain message format
            for msg in conversation_history:
                role = msg.get("role", "")
                content = msg.get("content", "") or msg.get("parts", [""])[0]  # Support both formats
                if role == "user":
                    messages.append(HumanMessage(content=content))
                elif role in ("assistant", "model"):
                    messages.append(AIMessage(content=content))
        else:
            # Initialize with the current NDA as context
            messages.append(HumanMessage(content="Please review this NDA."))
            messages.append(AIMessage(content=nda_text))
        
        # Add the modification request
        messages.append(HumanMessage(content=modification_request))
        
        logger.info(f"Modifying NDA with request: {modification_request[:100]}...")
        
        try:
            # Invoke the model
            response = self.model.invoke(messages)
            modified_text = response.content
            
            logger.info("NDA modified successfully")
            return modified_text
        except Exception as e:
            logger.error(f"Error modifying NDA: {e}")
            raise Exception(f"Failed to modify NDA: {str(e)}")
    
    def create_docx(self, nda_text: str, client_name: str = "", counterparty_name: str = "") -> bytes:
        """
        Create a Word document from NDA text.
        
        Args:
            nda_text: The generated NDA text.
            client_name: Optional client name for filename generation.
            counterparty_name: Optional counterparty name for filename generation.
        
        Returns:
            Word document as bytes.
        """
        doc = Document()
        doc.add_heading('Non-Disclosure Agreement', 0)
        
        # Simple parsing: add paragraphs based on newlines
        for para in nda_text.split('\n'):
            # Check if the line is likely a heading (e.g., "Article 1. Definitions")
            stripped = para.strip()
            if stripped.lower().startswith(("article", "section")) or stripped.endswith(":"):
                doc.add_heading(stripped, level=2)
            elif stripped:  # Avoid adding empty paragraphs
                doc.add_paragraph(para)
        
        # Save the document to a byte stream
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        
        logger.info("Word document created successfully")
        return bio.getvalue()
    
    def generate_nda_with_docx(self, user_inputs: dict) -> Tuple[str, bytes]:
        """
        Generate NDA text and create a Word document.
        
        Args:
            user_inputs: Dictionary containing NDA parameters.
        
        Returns:
            Tuple of (nda_text, docx_bytes).
        """
        nda_text = self.generate_nda_text(user_inputs)
        docx_bytes = self.create_docx(
            nda_text, 
            user_inputs.get("first_party", ""),
            user_inputs.get("second_party", "")
        )
        return nda_text, docx_bytes
