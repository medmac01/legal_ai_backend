"""
Description:
Naive document parser implementation. Loads and splits documents (PDF, DOCX, TXT) into chunks
using recursive character splitting with configurable chunk size and overlap.

Author: Raptopoulos Petros [petrosrapto@gmail.com]
Date : 2025/03/10
"""
from typing import List
import os

from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.vectorstores.utils import filter_complex_metadata
from langchain_unstructured import UnstructuredLoader

from docx import Document as DocxDocument

class NaiveDocumentSplitterAndParser:
    """Parser class that loads a file, splits its content, and returns a list of Document objects."""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Args:
            chunk_size: Maximum number of characters per chunk
            chunk_overlap: Overlapping characters between consecutive chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
        )

    def parse_document(self, file_path: str) -> List[Document]:
        """
        Loads and splits the file at file_path into smaller chunks.

        Args:
            file_path: Path to the file (PDF or text)
        
        Returns:
            A list of LangChain Document objects
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        # Identify file extension
        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension == ".pdf":
            # Use PyPDFLoader for PDFs
            # loader = PyPDFLoader(file_path)
            # docs = loader.load()
            loader = UnstructuredLoader(file_path=file_path, mode="elements", strategy="hi_res",)
            elements = list(loader.lazy_load())
            # Combine all text content
            full_text = "\n".join([el.page_content for el in elements])
            # Get metadata from the first element
            # base_metadata = filter_complex_metadata(elements[0].metadata if elements else {"source": file_path})
            # Create one merged Document with filtered metadata
            docs = [Document(page_content=full_text, metadata={"source": file_path})]
        elif file_extension in [".doc", ".docx"]:
            # Use python-docx to read Word documents
            doc = DocxDocument(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])  # Extract text
            docs = [Document(page_content=text, metadata={"source": file_path})]
        else:
            # Assume it's a text file if not PDF
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            # Create a single Document to hold text content
            docs = [Document(page_content=text, metadata={"source": file_path})]

        # for doc in docs:
        #     doc.metadata = filter_complex_metadata(doc.metadata)
        # Split into smaller chunks
        splitted_docs = self.splitter.split_documents(docs)
        
        # Add formatting to each chunk indicating it's part of the original document
        for doc in splitted_docs:
            doc.page_content = f"""--- ORIGINAL SPAN OF THE DOCUMENT ---\n{doc.page_content}\n------"""

        return splitted_docs